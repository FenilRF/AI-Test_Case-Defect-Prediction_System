"""
API Routes — Design Documentation Upload
-------------------------------------------
Endpoints:
  POST /api/design/upload           → upload design files/images/URLs
  GET  /api/design/{design_id}      → get design document details
  GET  /api/design/by-requirement/{req_id} → list designs for a requirement
"""

import json
import logging
import os
import shutil
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy.orm import Session

from app.models.database import get_db
from app.models.models import DesignDocument, Requirement
from app.schemas.schemas import (
    DesignUploadResponse,
    DesignAnalysisResult,
    DesignDocumentOut,
    TestCaseOut,
)
from app.services.design_analysis_service import analyse_design
from app.services.folder_storage_service import (
    create_design_folder,
    save_design_links,
)
from app.services.test_case_generator import generate_test_cases
from app.services.nlp_service import parse_requirement

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/design", tags=["Design Documentation"])

# Allowed file extensions
_ALLOWED_DOC_EXTENSIONS = {".pdf", ".docx", ".pptx", ".txt"}
_ALLOWED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
_ALLOWED_EXTENSIONS = _ALLOWED_DOC_EXTENSIONS | _ALLOWED_IMAGE_EXTENSIONS
_MAX_FILE_SIZE_MB = 20


def _validate_file(file: UploadFile) -> str:
    """Validate uploaded file extension. Returns the lowercase extension."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="File has no name")
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in _ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Allowed: {', '.join(sorted(_ALLOWED_EXTENSIONS))}",
        )
    return ext


def _extract_text_from_file(file_path: str, ext: str) -> str:
    """
    Extract text content from uploaded files for analysis.
    Only extracts from TXT files — other formats store filename as context.
    """
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.warning("Could not read text from %s: %s", file_path, e)
    elif ext == ".pdf":
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        except Exception as e:
            logger.warning("PDF extraction failed (PyPDF2 might not be installed): %s", e)
    elif ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            logger.warning("DOCX extraction failed (python-docx might not be installed): %s", e)
    # For non-text files, return the filename as minimal context
    return os.path.basename(file_path)


@router.post("/upload", response_model=DesignUploadResponse)
async def upload_design(
    design_text: str = Form(None),
    design_url: str = Form(None),
    file: UploadFile = File(None),
    requirement_id: Optional[int] = Form(None, description="Associated requirement ID"),
    db: Session = Depends(get_db),
):
    """
    Upload design documentation — files, images, and/or URLs.

    At least one input (file, image, or URL) is required.
    Associates design with requirement_id if provided.
    Stores design inside the requirement-specific folder under design/.

    After upload, performs lightweight design analysis.
    """
    # ── Parse URLs ────────────────────────────────────────
    import re
    design_urls = []
    if design_url and design_url.strip():
        u = design_url.strip()
        url_pattern = re.compile(r"^(https?://)([a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}(/.*)?$")
        if not u.startswith("http"):
            u = "https://" + u
        if not url_pattern.match(u):
            raise HTTPException(status_code=400, detail=f"Invalid URL format: {u}")
        design_urls.append(u)

    # ── Collect all upload files ──────────────────────────
    all_files: List[UploadFile] = []
    if file:
        all_files.append(file)

    # ── Validate: at least one input required ─────────────
    if not all_files and not design_urls and not design_text:
        raise HTTPException(
            status_code=400,
            detail="At least one input is required: text, file, image, or URL.",
        )

    # ── Validate requirement_id (if provided) ─────────────
    requirement_folder_path = None
    if requirement_id:
        req = db.query(Requirement).filter(Requirement.id == requirement_id).first()
        if not req:
            raise HTTPException(status_code=404, detail=f"Requirement #{requirement_id} not found")
        requirement_folder_path = req.folder_path

    # ── Create temp design folder ─────────────────────────
    design_folder = create_design_folder(module_name="Unknown")

    # ── Validate & save files ─────────────────────────────
    saved_file_names = []
    saved_file_paths = []
    extracted_texts = []

    for f in all_files:
        ext = _validate_file(f)

        # Safe filename
        safe_name = f.filename.replace("..", "_").replace("/", "_").replace("\\", "_")
        dest_path = os.path.join(design_folder, safe_name)

        # Save file
        try:
            contents = await f.read()
            if len(contents) > _MAX_FILE_SIZE_MB * 1024 * 1024:
                raise HTTPException(
                    status_code=400,
                    detail=f"File '{f.filename}' exceeds {_MAX_FILE_SIZE_MB}MB limit",
                )
            with open(dest_path, "wb") as out:
                out.write(contents)
            saved_file_names.append(safe_name)
            saved_file_paths.append(dest_path)
            logger.info("Saved design file: %s → %s", f.filename, dest_path)

            # Extract text for analysis
            text = _extract_text_from_file(dest_path, ext)
            if text:
                extracted_texts.append(text)
        except HTTPException:
            raise
        except Exception as e:
            logger.exception("Failed to save file %s: %s", f.filename, e)
            raise HTTPException(status_code=500, detail=f"Failed to save file '{f.filename}': {str(e)}")

    # ── Save design URLs ──────────────────────────────────
    if design_urls:
        save_design_links(design_folder, design_urls)
        # Add URL text to analysis context
        for u in design_urls:
            extracted_texts.append(f"design url: {u}")

    if design_text and design_text.strip():
        extracted_texts.append(design_text.strip())

    # ── Design analysis ───────────────────────────────────
    combined_text = "\n".join(extracted_texts) if extracted_texts else "Design documentation uploaded"
    analysis = analyse_design(
        combined_text=combined_text,
        file_names=saved_file_names,
        urls=design_urls,
    )

    # ── Update folder name with module ────────────────────
    new_folder = design_folder
    if analysis.get("modules"):
        mod_name = analysis["modules"][0]
        base_dir = os.path.dirname(design_folder)
        base_name = os.path.basename(design_folder)
        new_base = base_name.replace("Unknown", mod_name.replace(" ", "_"))
        new_folder = os.path.join(base_dir, new_base)
        if new_folder != design_folder:
            try:
                os.rename(design_folder, new_folder)
                
                # Update paths in saved_file_paths to point to new folder
                saved_file_paths = [p.replace(design_folder, new_folder) for p in saved_file_paths]
                logger.info("Renamed design folder to: %s", new_folder)
            except Exception as e:
                logger.warning("Could not rename design folder to %s: %s", new_folder, e)
                new_folder = design_folder

    # ── Generate Test Cases ───────────────────────────────
    try:
        parsed_req = parse_requirement(combined_text)
        raw_test_cases = generate_test_cases(parsed_req)

        # Ensure ID generation or use dummy IDs for design tests
        parsed_test_cases = []
        for i, tc in enumerate(raw_test_cases, 1):
            parsed_test_cases.append({
                "test_id": i + 10000, # Fake ID series or we insert to DB
                "module_name": tc.get("module_name", "Unknown"),
                "scenario": tc.get("scenario", ""),
                "test_type": tc.get("test_type", ""),
                "test_level": tc.get("test_level", "Unit"),
                "expected_result": tc.get("expected_result", ""),
                "priority": tc.get("priority", "P3"),
            })
            
        tc_file_path = os.path.join(new_folder, "test_cases.json")
        with open(tc_file_path, "w", encoding="utf-8") as f:
            json.dump(parsed_test_cases, f, indent=2)
        logger.info("Saved generated design test cases to %s", tc_file_path)
    except Exception as e:
        logger.warning("Test cases could not be generated for design: %s", e)
        parsed_test_cases = []

    # ── Store metadata in database ────────────────────────
    design_doc = DesignDocument(
        requirement_id=requirement_id,
        file_names=json.dumps(saved_file_names),
        file_paths=json.dumps(saved_file_paths),
        design_urls=json.dumps(design_urls),
        analysis_result=json.dumps(analysis),
        folder_path=new_folder,
    )
    db.add(design_doc)
    db.commit()
    db.refresh(design_doc)

    logger.info(
        "Design uploaded: id=%d, req_id=%s, files=%d, urls=%d",
        design_doc.id, requirement_id, len(saved_file_names), len(design_urls),
    )

    # ── Build response ────────────────────────────────────
    file_count = len(saved_file_names)
    url_count = len(design_urls)
    msg_parts = []
    if file_count:
        msg_parts.append(f"{file_count} file(s) uploaded")
    if url_count:
        msg_parts.append(f"{url_count} URL(s) linked")
    message = " and ".join(msg_parts) + " successfully."

    return DesignUploadResponse(
        design_id=design_doc.id,
        requirement_id=requirement_id,
        uploaded_files=saved_file_names,
        design_urls=design_urls,
        folder_path=new_folder,
        analysis=DesignAnalysisResult(**analysis),
        message=message,
        test_cases=parsed_test_cases,
    )


@router.get("/{design_id}")
def get_design_document(design_id: int, db: Session = Depends(get_db)):
    """Retrieve a single design document by ID."""
    doc = db.query(DesignDocument).filter(DesignDocument.id == design_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Design document not found")

    return {
        "id": doc.id,
        "requirement_id": doc.requirement_id,
        "file_names": json.loads(doc.file_names) if doc.file_names else [],
        "design_urls": json.loads(doc.design_urls) if doc.design_urls else [],
        "folder_path": doc.folder_path,
        "analysis": json.loads(doc.analysis_result) if doc.analysis_result else None,
        "created_at": doc.created_at,
    }


@router.get("/by-requirement/{req_id}")
def list_designs_for_requirement(req_id: int, db: Session = Depends(get_db)):
    """List all design documents for a specific requirement."""
    req = db.query(Requirement).filter(Requirement.id == req_id).first()
    if not req:
        raise HTTPException(status_code=404, detail="Requirement not found")

    docs = db.query(DesignDocument).filter(DesignDocument.requirement_id == req_id).all()
    return [
        {
            "id": doc.id,
            "requirement_id": doc.requirement_id,
            "file_names": json.loads(doc.file_names) if doc.file_names else [],
            "design_urls": json.loads(doc.design_urls) if doc.design_urls else [],
            "folder_path": doc.folder_path,
            "analysis": json.loads(doc.analysis_result) if doc.analysis_result else None,
            "created_at": doc.created_at,
        }
        for doc in docs
    ]
